"""Read and write real documents — .tex, .docx, .md/.txt, .pdf — without flattening them.

The naive way to "edit a document with a model" is to extract plain text, rewrite the
whole thing, and save the result. That destroys everything a document actually is:
LaTeX structure, Word styles, headings, tables, equations. It also makes the edit
unreviewable, because the output shares no structure with the input.

So this module works in **segments**. A document is read as an ordered list of editable
paragraphs plus the scaffolding between them. A rewrite touches one paragraph at a time,
each is independently accepted or rejected, and the document is written back through the
same structure it came from — a .docx stays a .docx with its styles, a .tex keeps its
preamble, macros and math.
"""
from __future__ import annotations

import re
import shutil
import subprocess
from dataclasses import dataclass, field
from pathlib import Path

SPIRAL_CLAY_HEX = "D97757"
SPIRAL_CLAY_RGB = (0xD9, 0x77, 0x57)

# LaTeX constructs whose contents must never be handed to a prose rewriter
_TEX_PROTECTED = re.compile(
    r"\\begin\{(equation|align|gather|multline|eqnarray|figure|table|tabular|"
    r"lstlisting|verbatim|algorithm|thebibliography)\*?\}.*?"
    r"\\end\{\1\*?\}|"
    r"\$\$.*?\$\$|"
    r"\\\[.*?\\\]|"
    # NOTE: [^\n]* not .* — this pattern is compiled with DOTALL for the environment
    # alternatives above, and a greedy .* here swallowed the entire document.
    r"^[ \t]*\\(?:documentclass|usepackage|newcommand|renewcommand|def|input|include|"
    r"bibliography|bibliographystyle)\b[^\n]*$",
    re.S | re.M)


@dataclass
class Segment:
    """One editable unit. ``editable`` is False for scaffolding that must pass through
    byte-identical (preamble, equations, figures, code)."""
    index: int
    text: str
    editable: bool = True
    kind: str = "paragraph"


@dataclass
class Document:
    path: Path
    kind: str                      # tex | docx | markdown | pdf
    segments: list = field(default_factory=list)
    _backing: object = None        # docx Document, when applicable

    @property
    def editable_segments(self) -> list:
        return [s for s in self.segments if s.editable and s.text.strip()]

    def text(self) -> str:
        return "\n\n".join(s.text for s in self.segments)

    def with_replacements(self, replacements: dict) -> str:
        """Full text with ``{index: new_text}`` applied — used for measuring the result
        before anything is written to disk."""
        if self.kind == "tex" and isinstance(self._backing, str):
            # TeX whitespace is usually harmless, until it is not: a blank line inside a
            # multiline \title{...}, \caption{...}, macro argument, or moving argument
            # creates a paragraph token and can make the project uncompilable. Splice
            # replacements into the exact raw source instead of rejoining segments.
            raw = self._backing
            cursor = 0
            chunks = []
            for segment in self.segments:
                position = raw.find(segment.text, cursor)
                if position < 0:
                    raise RuntimeError(
                        f"cannot locate TeX segment {segment.index} in original source"
                    )
                chunks.append(raw[cursor:position])
                chunks.append(replacements.get(segment.index, segment.text))
                cursor = position + len(segment.text)
            chunks.append(raw[cursor:])
            return "".join(chunks)
        out = []
        for s in self.segments:
            out.append(replacements.get(s.index, s.text))
        return "\n\n".join(out)


# ── reading ──────────────────────────────────────────────────────────────────
def _read_tex(path: Path) -> Document:
    raw = path.read_text(errors="replace")
    segments: list[Segment] = []
    cursor = 0
    idx = 0

    # a line that is nothing but a LaTeX command is structure, not prose — it must not
    # be handed to a rewriter even when it sits inside a paragraph block
    command_line = re.compile(
        r"^[ \t]*\\(?:begin|end|section|subsection|subsubsection|chapter|part|"
        r"paragraph|title|author|date|maketitle|label|caption|item|bibliography\w*|"
        r"appendix|tableofcontents|newpage|clearpage|noindent|centering)\b[^\n]*$|"
        r"^[ \t]*%[^\n]*$")

    def add_prose(chunk: str) -> None:
        nonlocal idx
        for para in re.split(r"\n\s*\n", chunk):
            if not para.strip():
                continue
            buf: list[str] = []

            def flush_buf() -> None:
                nonlocal idx, buf
                if buf and "".join(buf).strip():
                    segments.append(Segment(idx, "\n".join(buf), editable=True,
                                            kind="paragraph"))
                    idx += 1
                buf = []

            for line in para.split("\n"):
                if command_line.match(line):
                    flush_buf()
                    segments.append(Segment(idx, line, editable=False, kind="command"))
                    idx += 1
                else:
                    buf.append(line)
            flush_buf()

    for m in _TEX_PROTECTED.finditer(raw):
        add_prose(raw[cursor:m.start()])
        segments.append(Segment(idx, m.group(0), editable=False, kind="protected"))
        idx += 1
        cursor = m.end()
    add_prose(raw[cursor:])
    return Document(path, "tex", segments, _backing=raw)


def _read_docx(path: Path) -> Document:
    try:
        import docx
    except ImportError as exc:                       # pragma: no cover - env dependent
        raise RuntimeError(
            "reading .docx needs python-docx: pip install python-docx") from exc
    doc = docx.Document(str(path))
    segments: list[Segment] = []
    for i, para in enumerate(doc.paragraphs):
        style = (para.style.name if para.style is not None else "") or ""
        # headings and captions carry meaning in few words; rewriting them adds risk
        # without benefit, so they pass through untouched
        structural = bool(re.search(
            r"(?i)(?:^|\s)(?:heading|title|caption|toc|quote)(?:\s|$)", style))
        # python-docx exposes only part of Word's inline object model. Replacing a
        # paragraph that contains fields, hyperlinks, review markup, bookmarks, or
        # drawings can orphan relationships or silently erase reviewer intent. Such
        # paragraphs remain measurable but are never handed to the rewriter.
        complex_inline = any(para._p.xpath(f".//w:{tag}") for tag in (
            "hyperlink", "fldChar", "instrText", "bookmarkStart", "bookmarkEnd",
            "drawing", "object", "sdt", "commentRangeStart", "commentRangeEnd",
            "commentReference", "ins", "del",
        ))
        editable = not structural and not complex_inline
        segments.append(Segment(i, para.text, editable=editable,
                                kind=style.lower() or "paragraph"))
    return Document(path, "docx", segments, _backing=doc)


def _read_flat(path: Path, kind: str) -> Document:
    raw = (path.read_text(errors="replace") if kind != "pdf"
           else _read_pdf_text(path))
    segments = []
    for i, para in enumerate(re.split(r"\n\s*\n", raw)):
        heading = bool(re.match(r"\s*#{1,6}\s", para))
        segments.append(Segment(i, para, editable=not heading,
                                kind="heading" if heading else "paragraph"))
    return Document(path, kind, segments)


def _read_pdf_text(path: Path) -> str:
    from spiral.research_corpus import extract_pdf_text

    return extract_pdf_text(path)


def read_document(path: str | Path) -> Document:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".tex":
        return _read_tex(p)
    if suffix == ".docx":
        return _read_docx(p)
    if suffix == ".pdf":
        return _read_flat(p, "pdf")
    return _read_flat(p, "markdown")


def mark_tex_change(text: str) -> str:
    """Mark substantive TeX edits with Spiral's clay accent."""

    return "{\\color{SpiralClay}" + str(text) + "}"


def ensure_tex_change_color(text: str) -> str:
    """Declare the review color in a copied TeX document, never the source."""

    if r"\definecolor{SpiralClay}" in text:
        return text
    definition = rf"\definecolor{{SpiralClay}}{{HTML}}{{{SPIRAL_CLAY_HEX}}}"
    xcolor = re.search(r"(?m)^\s*\\usepackage(?:\[[^\]]*\])?\{xcolor\}\s*$", text)
    if xcolor:
        return text[:xcolor.end()] + "\n" + definition + text[xcolor.end():]
    documentclass = re.search(r"(?m)^\s*\\documentclass(?:\[[^\]]*\])?\{[^}]+\}\s*$",
                              text)
    declaration = "\\usepackage{xcolor}\n" + definition
    if documentclass:
        return text[:documentclass.end()] + "\n\n" + declaration + text[documentclass.end():]
    return declaration + "\n" + text


# ── writing ──────────────────────────────────────────────────────────────────
def write_document(doc: Document, replacements: dict, out_path: str | Path) -> Path:
    """Write the edited document, preserving its native structure.

    A .docx is written back through python-docx so styles, headings and everything the
    rewriter never touched survive; a .tex is reassembled with its preamble and math
    exactly as they were. A PDF cannot be faithfully rewritten, so its edit is emitted
    as markdown and that is stated rather than pretended otherwise."""
    out = require_distinct_output(doc.path, out_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    if doc.kind == "docx" and doc._backing is not None:
        from docx.shared import RGBColor

        backing = doc._backing
        for seg_index, new_text in replacements.items():
            if seg_index >= len(backing.paragraphs):
                continue
            para = backing.paragraphs[seg_index]
            if not para.runs:
                para.text = new_text
                if para.runs:
                    para.runs[0].font.color.rgb = RGBColor(*SPIRAL_CLAY_RGB)
                continue
            # keep the first run's formatting, drop the rest — the paragraph keeps its
            # font/style instead of reverting to document defaults
            para.runs[0].text = new_text
            para.runs[0].font.color.rgb = RGBColor(*SPIRAL_CLAY_RGB)
            for run in para.runs[1:]:
                run.text = ""
        backing.save(str(out))
        return out

    marked = replacements
    if doc.kind == "tex" and replacements:
        marked = {index: mark_tex_change(value) for index, value in replacements.items()}
        out.write_text(ensure_tex_change_color(doc.with_replacements(marked)))
        return out
    out.write_text(doc.with_replacements(marked))
    return out


def require_distinct_output(source: str | Path, output: str | Path) -> Path:
    """Reject every filesystem alias of the source, including symlinks/hard links.

    This is the last line of defence beneath the CLI.  A caller may choose any output
    name, but no rewrite operation is ever allowed to target the input inode.
    """

    source_path = Path(source)
    output_path = Path(output)
    if source_path.resolve() == output_path.resolve():
        raise ValueError("output must be a copy; it cannot be the source document")
    if output_path.exists():
        try:
            if source_path.samefile(output_path):
                raise ValueError(
                    "output must be a copy; it aliases the source document"
                )
        except FileNotFoundError:
            pass
    return output_path


_PROJECT_CACHE_DIRS = {
    ".git", ".lake", ".mypy_cache", ".pytest_cache", ".spiral", ".venv",
    "__pycache__", "node_modules", "venv",
}
_LATEX_BUILD_SUFFIXES = {
    ".aux", ".bbl", ".bcf", ".blg", ".fdb_latexmk", ".fls", ".log",
    ".out", ".run.xml", ".synctex.gz", ".toc",
}


def _unique_project_copy_root(source_root: Path) -> Path:
    base = source_root.with_name(source_root.name + "_spiral")
    candidate = base
    counter = 2
    while candidate.exists():
        candidate = base.with_name(f"{base.name}-{counter}")
        counter += 1
    return candidate


def prepare_tex_project_copy(source: str | Path) -> Path:
    """Copy a TeX source project and return the copied main-file path.

    Generated dependency/cache trees are omitted, but source code, bibliography files,
    figures, data, included TeX, and project documentation are copied.  The destination
    is always new, so an earlier edited project is never silently overwritten.
    """

    source = Path(source).resolve()
    source_root = source.parent
    destination = _unique_project_copy_root(source_root)

    def ignore(directory: str, names: list[str]) -> set[str]:
        directory_path = Path(directory)
        omitted = {name for name in names if name in _PROJECT_CACHE_DIRS}
        if directory_path.resolve() == source_root:
            for name in names:
                candidate = directory_path / name
                if candidate.is_file() and candidate.stem == source.stem:
                    suffix = "".join(candidate.suffixes[-2:])
                    if candidate.suffix in _LATEX_BUILD_SUFFIXES or suffix in _LATEX_BUILD_SUFFIXES:
                        omitted.add(name)
                if name == source.with_suffix(".pdf").name:
                    omitted.add(name)
        return omitted

    shutil.copytree(source_root, destination, ignore=ignore, copy_function=shutil.copy2)
    copied = destination / source.name
    if not copied.is_file():
        raise RuntimeError(f"TeX project copy is missing {source.name}")
    return copied


def compile_tex_copy(path: str | Path, *, timeout: int = 240) -> dict:
    """Compile a copied TeX main file without shell escape and return an audit record."""

    path = Path(path).resolve()
    latexmk = shutil.which("latexmk")
    if not latexmk:
        return {"available": False, "ok": False, "error": "latexmk is unavailable"}
    try:
        result = subprocess.run(
            [latexmk, "-pdf", "-interaction=nonstopmode", "-halt-on-error", path.name],
            cwd=path.parent, text=True, capture_output=True, timeout=timeout,
            check=False,
        )
    except Exception as exc:
        return {
            "available": True, "ok": False,
            "error": f"{type(exc).__name__}: {exc}",
        }
    combined = (result.stdout or "") + "\n" + (result.stderr or "")
    return {
        "available": True, "ok": result.returncode == 0,
        "returncode": result.returncode,
        "pdf": str(path.with_suffix(".pdf")) if path.with_suffix(".pdf").is_file() else "",
        "tail": combined[-5000:],
        "error": "" if result.returncode == 0 else "LaTeX compilation failed",
    }


def default_output_path(path: str | Path) -> Path:
    """Where an edit lands by default: alongside the original, never over it."""
    p = Path(path)
    if p.suffix.lower() == ".pdf":
        return p.with_name(f"{p.stem}_spiral.md")  # a PDF edit is emitted as markdown
    return p.with_name(f"{p.stem}_spiral{p.suffix or '.txt'}")
