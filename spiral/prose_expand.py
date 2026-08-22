"""Evidence expansion and structure alignment for ``spiral prose``.

These operations are intentionally separate from ordinary rewriting.  A line edit must
not invent facts; ``--beef-up`` is explicitly additive and therefore needs its own
evidence contract.  Likewise, ``--restructure`` moves intact section blocks instead of
asking a model to regenerate the document in a preferred order.

Corpus text is untrusted reference data.  The model sees bounded, identified anchors;
every retained sentence cites one of those identifiers, survives a second entailment
audit, and passes deterministic identifier, number, overlap, and exact-anchor checks.
"""
from __future__ import annotations

import hashlib
import html
import json
import re
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path


_WORD = re.compile(r"[A-Za-z][A-Za-z'-]+")
_NUMBER = re.compile(
    r"(?<![\w.])[+-]?(?:\d{1,3}(?:,\d{3})+|\d+|\.\d+)"
    r"(?:\.\d+)?(?:[eE][+-]?\d+)?%?(?!\w|\.\d)"
)
_SOURCE_MARKER = re.compile(r"\[(S\d+(?:\s*[;,]\s*S\d+)*)\]")
_ANCHOR_MARKER = re.compile(
    r"\[(S\d+-A\d+(?:\s*[;,]\s*S\d+-A\d+)*)\]"
)
_SENTENCE = re.compile(r"(?<=[.!?])\s+")
_MD_HEADING = re.compile(r"(?m)^(#{1,6})\s+([^\n#].*?)\s*$")
_TEX_SECTION = re.compile(r"(?m)^\s*\\section\*?\{([^{}]{1,160})\}\s*$")


@dataclass
class EvidenceAnchor:
    id: str
    text: str


@dataclass
class EvidenceSource:
    id: str
    paper_id: str
    title: str
    authors: list[str] = field(default_factory=list)
    year: str = "n.d."
    venue: str = ""
    doi: str = ""
    url: str = ""
    relevance: float = 0.0
    anchors: list[EvidenceAnchor] = field(default_factory=list)


@dataclass
class Addition:
    target_heading: str
    text: str
    evidence: list[dict]


@dataclass
class ExpansionResult:
    additions: list[Addition] = field(default_factory=list)
    sources: list[EvidenceSource] = field(default_factory=list)
    issues: list[str] = field(default_factory=list)

    @property
    def added_words(self) -> int:
        return sum(len(_WORD.findall(item.text)) for item in self.additions)


def _clean_source_text(raw: str) -> str:
    value = html.unescape(raw or "")
    value = re.sub(r"<[^>]+>", " ", value)
    value = re.sub(r"\\(?:sub)*section\*?\{([^{}]*)\}", r". \1. ", value)
    value = re.sub(r"\\(?:cite\w*|ref|label)\{[^{}]*\}", " ", value)
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def _content_terms(text: str, limit: int = 28) -> list[str]:
    try:
        from spiral.research_quality import topic_terms

        return topic_terms(text, limit=limit)
    except Exception:
        counts = Counter(word.lower() for word in _WORD.findall(text or "")
                         if len(word) >= 5)
        return [word for word, _ in counts.most_common(limit)]


def _anchor_candidates(paper, topic_terms: set[str]) -> list[tuple[float, str]]:
    abstract = _clean_source_text(getattr(paper, "abstract", "") or "")
    body = _clean_source_text(getattr(paper, "text", "") or "")
    pieces: list[tuple[str, bool, int]] = []
    for source, is_abstract in ((abstract, True), (body[:30_000], False)):
        pieces.extend((sentence, is_abstract, position)
                      for position, sentence in enumerate(_SENTENCE.split(source)))
    rows = []
    seen = set()
    for raw, is_abstract, position in pieces:
        sentence = " ".join(raw.split()).strip(" -")
        words = _WORD.findall(sentence)
        if not 9 <= len(words) <= 95:
            continue
        normalised = sentence.casefold()
        if (normalised in seen or "references" in normalised[:30]
                or re.search(r"\b(?:in|see|as discussed in) section\s+\d", normalised)
                or re.search(r"\bleft to future|come back to|begin in section\b",
                             normalised)):
            continue
        seen.add(normalised)
        held = {word.lower() for word in words}
        overlap = len(held & topic_terms)
        # Abstract/result anchors are preferred, but a concrete body sentence remains
        # available when the provider supplied only a short abstract.
        concrete = bool(_NUMBER.search(sentence) or re.search(
            r"\b(?:we (?:classify|compute|demonstrate|derive|find|found|generalize|"
            r"investigate|observe|observed|obtain|propose|show|study|measured|estimated)|"
            r"results?|associated|increased|decreased|higher|lower|compared|effect)\b",
            sentence, re.I))
        deictic = bool(re.match(
            r"(?i)(?:this|these|those|the dependence of this|similar structures|"
            r"we observe similar|as a result)\b", sentence,
        ))
        # Abstract claims summarize a paper's actual contribution and are far safer
        # additions than high-overlap TeX equation fragments or PDF section furniture.
        score = (overlap * 4 + (7 if concrete else 0) + (100 if is_abstract else 0)
                 - (30 if deictic else 0)
                 + (2 if position < 12 else 0))
        if overlap >= 1 or concrete:
            rows.append((float(score), sentence))
    rows.sort(key=lambda row: (-row[0], len(row[1]), row[1]))
    return rows


def build_evidence_packet(papers, selected_rows, draft: str, *, max_sources: int = 8,
                          anchors_per_source: int = 5,
                          minimum_sources: int = 4) -> list[EvidenceSource]:
    """Build bounded evidence from only the most relevant selected primary papers."""

    paper_list = list(papers or [])
    row_by_id = {str(row.get("id")): row for row in (selected_rows or [])
                 if isinstance(row, dict) and row.get("id")}
    scored = []
    for paper in paper_list:
        row = row_by_id.get(str(getattr(paper, "bare_id", "")), {})
        relevance = float(row.get("relevance") or 0.0)
        scored.append((relevance, paper, row))
    if not scored:
        return []
    # ``build_deep_profile`` has already enforced genre, population, full-text, topic,
    # and terminology gates.  Do not collapse that audited corpus to the top two merely
    # because their lexical scores happen to be higher than the rest.
    cutoff = 0.20
    shortlisted = [row for row in scored if row[0] >= cutoff]
    shortlisted.sort(key=lambda row: (-row[0], str(getattr(row[1], "bare_id", ""))))
    topic = set(_content_terms(draft, limit=32))
    sources: list[EvidenceSource] = []
    for source_number, (relevance, paper, _row) in enumerate(
            shortlisted[:max_sources], 1):
        candidates = _anchor_candidates(paper, topic)
        anchors = [
            EvidenceAnchor(f"S{source_number}-A{index}", text)
            for index, (_score, text) in enumerate(candidates[:anchors_per_source], 1)
        ]
        if not anchors:
            continue
        published = str(getattr(paper, "published", "") or "")
        match = re.search(r"(?:19|20)\d{2}", published)
        doi = str(getattr(paper, "doi", "") or "")
        url = str(getattr(paper, "url", "") or "")
        if not url and doi:
            url = "https://doi.org/" + doi
        sources.append(EvidenceSource(
            id=f"S{source_number}", paper_id=str(getattr(paper, "bare_id", "")),
            title=str(getattr(paper, "title", "") or "Untitled"),
            authors=list(getattr(paper, "authors", None) or []),
            year=match.group(0) if match else "n.d.",
            venue=str(getattr(paper, "venue", "") or ""), doi=doi, url=url,
            relevance=relevance, anchors=anchors,
        ))
    # Renumber after dropping sources without usable anchors.
    for source_number, source in enumerate(sources, 1):
        old = source.id
        source.id = f"S{source_number}"
        for anchor_number, anchor in enumerate(source.anchors, 1):
            anchor.id = f"{source.id}-A{anchor_number}"
        if old == source.id:
            continue
    if len(sources) < max(1, minimum_sources):
        return []
    return sources


def evidence_prompt(sources: list[EvidenceSource]) -> str:
    rows = []
    for source in sources:
        rows.append(
            f"SOURCE {source.id}\nTitle: {source.title}\n"
            f"Year: {source.year}\nRelevance: {source.relevance:.4f}"
        )
        rows.extend(f"- {anchor.id}: {anchor.text}" for anchor in source.anchors)
    return "\n\n".join(rows)


def document_headings(text: str, kind: str) -> list[str]:
    if kind == "tex":
        return [" ".join(match.split()) for match in _TEX_SECTION.findall(text or "")]
    return [" ".join(name.split()) for _marks, name in _MD_HEADING.findall(text or "")]


def _json_object(raw: str) -> dict:
    text = (raw or "").strip()
    text = re.sub(r"^```(?:json)?\s*|\s*```$", "", text, flags=re.I)
    try:
        value = json.loads(text)
        return value if isinstance(value, dict) else {}
    except Exception:
        return {}


def _marker_ids(text: str) -> list[str]:
    ids = []
    for group in _SOURCE_MARKER.findall(text or ""):
        ids.extend(re.findall(r"S\d+", group))
    return ids


def _used_source_ids(additions: list[Addition]) -> set[str]:
    return set(_marker_ids("\n".join(item.text for item in additions)))


def validate_additions(payload: dict, sources: list[EvidenceSource], original: str,
                       headings: list[str], *, max_added_words: int) -> ExpansionResult:
    """Deterministically validate a model proposal before semantic auditing."""

    source_by_id = {source.id: source for source in sources}
    anchor_by_id = {
        anchor.id: (source.id, anchor.text)
        for source in sources for anchor in source.anchors
    }
    allowed_headings = {" ".join(heading.split()).casefold(): heading for heading in headings}
    additions: list[Addition] = []
    issues: list[str] = []
    raw_additions = payload.get("additions") if isinstance(payload, dict) else None
    if not isinstance(raw_additions, list) or not raw_additions:
        return ExpansionResult(issues=["model returned no additions"])
    total_words = 0
    for index, raw in enumerate(raw_additions[:6], 1):
        if not isinstance(raw, dict):
            issues.append(f"addition {index} is not an object")
            continue
        text = " ".join(str(raw.get("text") or "").split())
        target = " ".join(str(raw.get("target_heading") or "").split())
        evidence = list(raw.get("evidence") or [])

        # Small local models reliably point at the exact anchor in the sentence itself
        # (``[S1-A2]``), but often omit a duplicated evidence array. Accept that stricter
        # form and derive the ordinary source citation mechanically. Unknown anchor IDs
        # stay in the text and fail the source-marker gate below.
        def anchor_marker(match) -> str:
            anchor_ids = re.findall(r"S\d+-A\d+", match.group(1))
            if any(anchor_id not in anchor_by_id for anchor_id in anchor_ids):
                return match.group(0)
            source_ids = []
            for anchor_id in anchor_ids:
                source_id = anchor_by_id[anchor_id][0]
                evidence.append({"source": source_id, "anchor": anchor_id})
                if source_id not in source_ids:
                    source_ids.append(source_id)
            return "[" + "; ".join(source_ids) + "]"

        text = _ANCHOR_MARKER.sub(anchor_marker, text)
        if len(_WORD.findall(text)) < 8:
            issues.append(f"addition {index} is too short")
            continue
        if target and target.casefold() not in allowed_headings:
            issues.append(f"addition {index} targets an unknown heading: {target}")
            continue
        target = allowed_headings.get(target.casefold(), target)
        declared: dict[str, set[str]] = {}
        for row in evidence if isinstance(evidence, list) else []:
            if not isinstance(row, dict):
                continue
            source_id = str(row.get("source") or "")
            anchor_id = str(row.get("anchor") or row.get("anchor_id") or "")
            held = anchor_by_id.get(anchor_id)
            if source_id in source_by_id and held and held[0] == source_id:
                declared.setdefault(source_id, set()).add(anchor_id)
        cited = _marker_ids(text)
        if not cited:
            issues.append(f"addition {index} has no source marker")
            continue
        if any(source_id not in source_by_id for source_id in cited):
            issues.append(f"addition {index} cites an unknown source")
            continue
        if any(not declared.get(source_id) for source_id in cited):
            issues.append(f"addition {index} lacks an exact anchor for a cited source")
            continue
        sentence_issue = False
        for sentence in _SENTENCE.split(text):
            if len(_WORD.findall(sentence)) < 5:
                continue
            sentence_sources = set(_marker_ids(sentence))
            if not sentence_sources:
                issues.append(f"addition {index} contains an uncited sentence")
                sentence_issue = True
                break
            anchors = [anchor_by_id[anchor_id][1]
                       for source_id in sentence_sources
                       for anchor_id in declared.get(source_id, set())]
            evidence_words = {word.lower() for anchor in anchors
                              for word in _content_terms(anchor, limit=50)}
            claim_words = set(_content_terms(_SOURCE_MARKER.sub("", sentence), limit=30))
            if len(claim_words & evidence_words) < 2:
                issues.append(
                    f"addition {index} has insufficient claim/evidence overlap; reuse at "
                    "least two distinctive content terms verbatim from the cited anchor "
                    "in each factual sentence"
                )
                sentence_issue = True
                break
            anchor_numbers = Counter(_NUMBER.findall(" ".join(anchors)))
            claim_numbers = Counter(_NUMBER.findall(_SOURCE_MARKER.sub("", sentence)))
            if claim_numbers - anchor_numbers:
                issues.append(f"addition {index} introduces a number absent from its anchors")
                sentence_issue = True
                break
        if sentence_issue:
            continue
        words = len(_WORD.findall(text))
        total_words += words
        if total_words > max_added_words:
            issues.append(
                f"additions exceed the {max_added_words}-word evidence expansion budget"
            )
            break
        additions.append(Addition(target, text, list(evidence)))
    if not additions and not issues:
        issues.append("no addition passed deterministic grounding checks")
    return ExpansionResult(additions=additions, sources=sources, issues=issues)


def _claim_rows(additions: list[Addition]) -> list[dict]:
    rows = []
    for addition_index, addition in enumerate(additions):
        declared = {}
        for row in addition.evidence:
            if isinstance(row, dict):
                declared.setdefault(str(row.get("source") or ""), []).append(
                    str(row.get("anchor") or row.get("anchor_id") or ""))
        for sentence in _SENTENCE.split(addition.text):
            if len(_WORD.findall(sentence)) < 5:
                continue
            rows.append({
                "claim_id": f"C{len(rows) + 1}", "text": sentence,
                "sources": sorted(set(_marker_ids(sentence))),
                "allowed_anchors": sorted({anchor for source in _marker_ids(sentence)
                                           for anchor in declared.get(source, [])}),
                "addition_index": addition_index,
            })
    return rows


def validate_entailment_audit(audit: dict, claims: list[dict]) -> list[str]:
    records = {str(row.get("claim_id")): row for row in audit.get("claims", [])
               if isinstance(row, dict) and row.get("claim_id")}
    issues = []
    for claim in claims:
        record = records.get(claim["claim_id"])
        if not record:
            issues.append(f"claim {claim['claim_id']} was not audited")
            continue
        if record.get("supported") is not True:
            issues.append(f"claim {claim['claim_id']} is not entailed by its anchors")
            continue
        anchor_ids = {str(value) for value in record.get("anchor_ids", [])}
        if not anchor_ids or not anchor_ids.issubset(set(claim["allowed_anchors"])):
            issues.append(f"claim {claim['claim_id']} used an unverified anchor")
    return issues


def generate_grounded_additions(ol, cfg, model: str, draft: str, kind: str, papers,
                                  selected_rows, *, rounds: int = 3,
                                  minimum_sources: int = 4) -> ExpansionResult:
    sources = build_evidence_packet(
        papers, selected_rows, draft, minimum_sources=minimum_sources,
    )
    if not sources:
        return ExpansionResult(issues=[
            f"fewer than {minimum_sources} closely matched primary sources had usable "
            "evidence anchors"
        ])
    accepted_source_ids = [source.id for source in sources]
    all_headings = document_headings(draft, kind)
    headings = [
        heading for heading in all_headings
        if _section_role(heading) in {"introduction", "setup", "discussion"}
    ]
    original_words = max(1, len(_WORD.findall(draft)))
    max_added_words = min(800, max(120, int(original_words * 0.40)))
    system = (
        "You are an evidence-bound academic editor. Add useful detail to the supplied "
        "article, but do not rewrite or restate its existing prose. Use only the exact "
        "EVIDENCE ANCHORS. Every added factual sentence must end with one or more exact "
        "anchor markers such as [S1-A1] or [S1-A1; S2-A3]. The only valid IDs are the "
        "ones printed below. Do not infer causality, generality, population, "
        "or numbers beyond the anchor. Return JSON only: {\"additions\":[{\"target_heading\":"
        "\"an exact allowed heading or empty\",\"text\":\"... [S1-A1].\"}]}. "
        "Do not target an abstract, methods/results section, or conclusion. Use direct, "
        "restrained academic prose without significance inflation, canned transitions, "
        "or promotional wording. Do not emit a separate evidence field. Return 4-6 "
        "additions, with EXACTLY ONE factual sentence in each addition, so every claim "
        "can pass or fail independently. In every sentence, reuse at least two "
        "distinctive content terms verbatim from its cited evidence anchor so the "
        "deterministic grounding "
        "check can verify the link. "
        "Each sentence must stand alone: do not begin with This, These, Similar, It, "
        "or another backward reference. "
        f"You MUST cite at least {minimum_sources} distinct ACCEPTED SOURCE "
        "IDs printed by the user. Include one compact, independently supported sentence "
        "for every source you use."
    )
    prompt = (
        f"ADDITION WORD BUDGET: {max_added_words}\n"
        f"ACCEPTED SOURCE IDS: {json.dumps(accepted_source_ids)}\n"
        f"MINIMUM DISTINCT SOURCES: {minimum_sources}\n"
        f"ALLOWED TARGET HEADINGS: {json.dumps(headings, ensure_ascii=False)}\n\n"
        f"ARTICLE:\n{draft[:20_000]}\n\nEVIDENCE ANCHORS:\n{evidence_prompt(sources)}"
    )
    feedback = []
    result = ExpansionResult(sources=sources)
    for _attempt in range(max(1, rounds)):
        user = prompt
        if feedback:
            user += "\n\nFIX THESE REJECTIONS:\n- " + "\n- ".join(feedback[-6:])
        response = ol.chat(
            model, [{"role": "system", "content": system},
                    {"role": "user", "content": user}],
            fmt="json", num_predict=4096, temperature=0.15,
            num_ctx=cfg.spec_for(model).num_ctx, keep_alive=cfg.keep_alive,
        )
        candidate = validate_additions(
            _json_object(getattr(response, "text", "")), sources, draft, headings,
            max_added_words=max_added_words,
        )
        if not candidate.additions:
            feedback.extend(candidate.issues)
            result.issues = list(dict.fromkeys(feedback))
            continue
        if kind == "tex":
            unsafe = sorted({character for addition in candidate.additions
                             for character in addition.text if ord(character) > 127})
            if unsafe:
                feedback.append(
                    "LaTeX addition contains non-ASCII source-extraction glyphs "
                    f"({''.join(unsafe[:8])}); express notation with explicit LaTeX "
                    "commands in math mode and remove extraction artifacts"
                )
                continue
        used_sources = _used_source_ids(candidate.additions)
        if len(used_sources) < minimum_sources:
            feedback.append(
                f"proposal used only {len(used_sources)} distinct accepted source(s); "
                f"use at least {minimum_sources} distinct sources"
            )
            continue
        try:
            from spiral.research_writer import suspicious_phrase_overlap

            overlap = suspicious_phrase_overlap(
                "\n\n".join(item.text for item in candidate.additions), papers, words=14,
            )
        except Exception as exc:
            feedback.append(f"source-overlap check failed: {type(exc).__name__}: {exc}")
            continue
        if overlap:
            feedback.append(
                "addition copies 14 source words; paraphrase while keeping the citation"
            )
            continue
        claims = _claim_rows(candidate.additions)
        audit_sources = {
            anchor.id: anchor.text for source in sources for anchor in source.anchors
        }
        audit_system = (
            "Audit whether each claim is fully entailed by its allowed exact anchors. "
            "Reject stronger causality, certainty, scope, population, or comparison. "
            "Return JSON only: {\"claims\":[{\"claim_id\":\"C1\","
            "\"supported\":true,\"anchor_ids\":[\"S1-A1\"]}]}."
        )
        audit_prompt = json.dumps({
            "claims": claims,
            "anchors": {anchor_id: audit_sources[anchor_id]
                        for claim in claims for anchor_id in claim["allowed_anchors"]},
        }, ensure_ascii=False)
        audit_model = cfg.critic.name or model
        audit_response = ol.chat(
            audit_model,
            [{"role": "system", "content": audit_system},
             {"role": "user", "content": audit_prompt}],
            fmt="json", num_predict=3072, temperature=0.0,
            num_ctx=cfg.spec_for(audit_model).num_ctx, keep_alive=cfg.keep_alive,
        )
        audit_issues = validate_entailment_audit(
            _json_object(getattr(audit_response, "text", "")), claims,
        )
        if audit_issues:
            audit_payload = _json_object(getattr(audit_response, "text", ""))
            safe_additions = set(range(len(candidate.additions)))
            for claim in claims:
                if validate_entailment_audit(audit_payload, [claim]):
                    safe_additions.discard(int(claim["addition_index"]))
            if safe_additions:
                candidate.additions = [
                    addition for index, addition in enumerate(candidate.additions)
                    if index in safe_additions
                ]
                surviving_sources = _used_source_ids(candidate.additions)
                if len(surviving_sources) < minimum_sources:
                    feedback.extend(audit_issues)
                    feedback.append(
                        "after entailment filtering, fewer than "
                        f"{minimum_sources} distinct accepted sources remained"
                    )
                    continue
                candidate.issues = audit_issues
                candidate.sources = [
                    source for source in sources
                    if source.id in _used_source_ids(candidate.additions)
                ]
                return candidate
            feedback.extend(audit_issues)
            continue
        candidate.issues = []
        candidate.sources = [
            source for source in sources
            if source.id in _used_source_ids(candidate.additions)
        ]
        return candidate

    # A single multi-source proposal is efficient, but smaller local models can satisfy
    # the prose task while repeatedly omitting one citation or mixing a weak sentence
    # into an otherwise grounded paragraph. Fall back to one independently gated claim
    # per source. This does not relax any acceptance rule: the combined result must still
    # contain the requested number of distinct sources and pass overlap + entailment.
    individual_system = (
        "Write exactly one compact factual sentence that adds useful academic context "
        "to the article. Use only ONE of the supplied exact evidence anchors. Reuse at "
        "least two distinctive content terms verbatim from that anchor, without copying "
        "14 consecutive words. Do not strengthen causality, certainty, scope, or numbers. "
        "End the sentence with its exact anchor marker. Return JSON only: "
        "{\"additions\":[{\"target_heading\":\"an exact allowed heading or empty\","
        "\"text\":\"one sentence [S1-A1].\"}]}. Do not use non-ASCII notation in "
        "LaTeX; express notation with explicit LaTeX commands in math mode."
        " The sentence must stand alone and must not begin with This, These, Similar, "
        "It, or another backward reference."
    )
    individual_additions: list[Addition] = []
    for source in sources:
        source_rejections: list[str] = []
        individual_prompt = (
            f"ONLY ACCEPTED SOURCE ID: {source.id}\n"
            f"ALLOWED TARGET HEADINGS: {json.dumps(headings, ensure_ascii=False)}\n\n"
            f"ARTICLE CONTEXT:\n{draft[:10_000]}\n\n"
            f"EVIDENCE ANCHORS:\n{evidence_prompt([source])}"
        )
        for _source_attempt in range(2):
            user = individual_prompt
            if source_rejections:
                user += "\n\nFIX THIS REJECTION:\n- " + source_rejections[-1]
            try:
                response = ol.chat(
                    model, [{"role": "system", "content": individual_system},
                            {"role": "user", "content": user}],
                    fmt="json", num_predict=1024, temperature=0.1,
                    num_ctx=cfg.spec_for(model).num_ctx, keep_alive=cfg.keep_alive,
                )
            except Exception as exc:
                source_rejections.append(
                    f"source {source.id} generation failed: {type(exc).__name__}: {exc}"
                )
                break
            held = validate_additions(
                _json_object(getattr(response, "text", "")), [source], draft, headings,
                max_added_words=min(140, max_added_words),
            )
            if not held.additions:
                source_rejections.extend(held.issues)
                continue
            addition = held.additions[0]
            if kind == "tex" and any(ord(character) > 127 for character in addition.text):
                source_rejections.append(
                    "LaTeX addition contains non-ASCII source-extraction glyphs"
                )
                continue
            try:
                from spiral.research_writer import suspicious_phrase_overlap

                overlap = suspicious_phrase_overlap(addition.text, papers, words=14)
            except Exception as exc:
                source_rejections.append(
                    f"source-overlap check failed: {type(exc).__name__}: {exc}"
                )
                continue
            if overlap:
                source_rejections.append(
                    "addition copies 14 source words; paraphrase while keeping the citation"
                )
                continue
            individual_additions.append(addition)
            break
        feedback.extend(source_rejections)

    if len(_used_source_ids(individual_additions)) >= minimum_sources:
        candidate = ExpansionResult(additions=individual_additions, sources=sources)
        claims = _claim_rows(candidate.additions)
        audit_sources = {
            anchor.id: anchor.text for source in sources for anchor in source.anchors
        }
        audit_prompt = json.dumps({
            "claims": claims,
            "anchors": {anchor_id: audit_sources[anchor_id]
                        for claim in claims for anchor_id in claim["allowed_anchors"]},
        }, ensure_ascii=False)
        audit_model = cfg.critic.name or model
        try:
            audit_response = ol.chat(
                audit_model,
                [{"role": "system", "content": (
                    "Audit whether each claim is fully entailed by its allowed exact "
                    "anchors. Reject stronger causality, certainty, scope, population, "
                    "or comparison. Return JSON only: {\"claims\":[{\"claim_id\":"
                    "\"C1\",\"supported\":true,\"anchor_ids\":[\"S1-A1\"]}]}.")},
                 {"role": "user", "content": audit_prompt}],
                fmt="json", num_predict=3072, temperature=0.0,
                num_ctx=cfg.spec_for(audit_model).num_ctx, keep_alive=cfg.keep_alive,
            )
            audit_payload = _json_object(getattr(audit_response, "text", ""))
            safe_additions = set(range(len(candidate.additions)))
            for claim in claims:
                if validate_entailment_audit(audit_payload, [claim]):
                    safe_additions.discard(int(claim["addition_index"]))
            candidate.additions = [
                addition for index, addition in enumerate(candidate.additions)
                if index in safe_additions
            ]
            if len(_used_source_ids(candidate.additions)) >= minimum_sources:
                candidate.sources = [
                    source for source in sources
                    if source.id in _used_source_ids(candidate.additions)
                ]
                candidate.issues = validate_entailment_audit(audit_payload, claims)
                return candidate
            feedback.append(
                "source-by-source fallback retained fewer than "
                f"{minimum_sources} sources after entailment auditing"
            )
        except Exception as exc:
            feedback.append(
                f"source-by-source entailment audit failed: {type(exc).__name__}: {exc}"
            )
    result.issues = list(dict.fromkeys(feedback)) or ["no grounded addition passed"]
    return result


def _short_author(source: EvidenceSource) -> str:
    if not source.authors:
        words = _WORD.findall(source.title)
        return " ".join(words[:3]) or "Corpus source"
    parts = source.authors[0].replace(",", " ").split()
    last = parts[-1]
    # PubMed commonly stores ``Surname AB`` while Crossref commonly stores
    # ``Given Surname``. Initial-only final tokens identify the former shape.
    first = (parts[0] if len(parts) > 1 and re.fullmatch(r"(?:[A-Z]\.?){1,5}", last)
             else last)
    return first + (" et al." if len(source.authors) > 1 else "")


def _citation_label(source: EvidenceSource) -> str:
    return f"{_short_author(source)}, {source.year}"


def render_citations(text: str, sources: list[EvidenceSource]) -> str:
    source_by_id = {source.id: source for source in sources}

    def replace(match) -> str:
        ids = re.findall(r"S\d+", match.group(1))
        labels = [_citation_label(source_by_id[source_id]) for source_id in ids
                  if source_id in source_by_id]
        return "(" + "; ".join(labels) + ")"

    return _SOURCE_MARKER.sub(replace, text)


def reference_entry(source: EvidenceSource) -> str:
    authors = ", ".join(source.authors) if source.authors else _short_author(source)
    venue = source.venue.strip(" .")
    locator = (" https://doi.org/" + source.doi if source.doi
               else (" " + source.url if source.url else ""))
    title = source.title.rstrip(" .")
    venue_part = f" {venue}." if venue else ""
    return f"{authors} ({source.year}). {title}.{venue_part}{locator}".strip()


def _normalise_doi(value: str) -> str:
    doi = str(value or "").strip().lower()
    doi = re.sub(r"^https?://(?:dx\.)?doi\.org/", "", doi)
    return doi.rstrip(" .")


def _normalise_title(value: str) -> str:
    return " ".join(re.findall(r"[a-z0-9]+", str(value or "").casefold()))


def _bibliography_files(tex_path: Path, text: str) -> list[Path]:
    names: list[str] = []
    for group in re.findall(r"\\bibliography\{([^{}]+)\}", text):
        names.extend(part.strip() for part in group.split(",") if part.strip())
    names.extend(re.findall(r"\\addbibresource(?:\[[^\]]*\])?\{([^{}]+)\}", text))
    paths = []
    for name in names:
        candidate = tex_path.parent / name
        if candidate.suffix.lower() != ".bib":
            candidate = Path(str(candidate) + ".bib")
        if candidate not in paths:
            paths.append(candidate)
    return paths


def _bibtex_index(paths: list[Path]) -> dict:
    """Index held BibTeX without rewriting the user's bibliography database."""

    keys: set[str] = set()
    doi_to_key: dict[str, str] = {}
    title_to_key: dict[str, str] = {}
    eprint_to_key: dict[str, str] = {}
    entry_start = re.compile(r"@\w+\s*\{\s*([^,\s]+)\s*,", re.I)
    field = lambda name, block: re.search(  # noqa: E731 - compact local parser
        rf"(?ims)^\s*{name}\s*=\s*[{{\"]([^}}\"]+)", block,
    )
    for path in paths:
        try:
            raw = path.read_text(encoding="utf-8", errors="replace")
        except OSError:
            continue
        starts = list(entry_start.finditer(raw))
        for index, match in enumerate(starts):
            key = match.group(1).strip()
            block = raw[match.start(): (starts[index + 1].start()
                                        if index + 1 < len(starts) else len(raw))]
            keys.add(key)
            doi_match = field("doi", block)
            title_match = field("title", block)
            eprint_match = field("eprint", block)
            if doi_match and _normalise_doi(doi_match.group(1)):
                doi_to_key.setdefault(_normalise_doi(doi_match.group(1)), key)
            if title_match and _normalise_title(title_match.group(1)):
                title_to_key.setdefault(_normalise_title(title_match.group(1)), key)
            if eprint_match:
                eprint_to_key.setdefault(eprint_match.group(1).strip().lower(), key)
    return {
        "keys": keys, "doi": doi_to_key, "title": title_to_key,
        "eprint": eprint_to_key,
    }


def _source_year(source: EvidenceSource) -> str:
    for value in (source.year, source.doi, source.paper_id):
        match = re.search(r"(?:19|20)\d{2}", str(value or ""))
        if match:
            return match.group(0)
    return ""


def _bibtex_author(name: str) -> str:
    clean = " ".join(str(name or "").split())
    if not clean or "," in clean:
        return clean
    parts = clean.split()
    if len(parts) > 1 and re.fullmatch(r"(?:[A-Z]\.?){1,6}", parts[-1]):
        return f"{parts[0]}, {' '.join(parts[1:])}"
    return clean


def _bibtex_escape(value: str) -> str:
    return re.sub(r"(?<!\\)([&%#_])", r"\\\1", str(value or ""))


def _generated_citation_key(source: EvidenceSource, held: set[str]) -> str:
    author = re.sub(r"[^A-Za-z0-9]", "", _short_author(source).split()[0]) or "Source"
    year = _source_year(source) or "ND"
    title_words = re.findall(r"[A-Za-z0-9]+", source.title)[:3]
    title = "".join(word[:1].upper() + word[1:] for word in title_words) or "Paper"
    identity = source.doi or source.paper_id or source.title
    digest = hashlib.sha256(identity.encode("utf-8", "ignore")).hexdigest()[:6]
    base = f"Spiral{author}{year}{title}{digest}"
    key = base
    counter = 2
    while key in held:
        key = f"{base}{counter}"
        counter += 1
    held.add(key)
    return key


def _bibtex_entry(source: EvidenceSource, key: str) -> str:
    fields = []
    authors = " and ".join(
        _bibtex_author(author) for author in source.authors if str(author).strip()
    )
    if authors:
        fields.append(("author", authors))
    fields.append(("title", "{" + _bibtex_escape(source.title) + "}"))
    year = _source_year(source)
    if year:
        fields.append(("year", year))
    if source.venue:
        fields.append(("journal", _bibtex_escape(source.venue)))
    if source.doi:
        fields.append(("doi", _normalise_doi(source.doi)))
    if source.url:
        fields.append(("url", _bibtex_escape(source.url)))
    body = ",\n".join(f"  {name} = {{{value}}}" for name, value in fields)
    return f"@article{{{key},\n{body}\n}}"


def _citation_plan(tex_path: Path, text: str, sources: list[EvidenceSource]) -> tuple[dict, list]:
    index = _bibtex_index(_bibliography_files(tex_path, text))
    held = set(index["keys"])
    keys: dict[str, str] = {}
    generated = []
    for source in sources:
        doi = _normalise_doi(source.doi or (
            source.paper_id.split(":", 1)[1]
            if source.paper_id.lower().startswith("doi:") else ""))
        title = _normalise_title(source.title)
        arxiv = source.paper_id.replace("arXiv:", "").lower()
        key = (index["doi"].get(doi) if doi else None)
        key = key or (index["title"].get(title) if title else None)
        key = key or index["eprint"].get(arxiv)
        if not key:
            key = _generated_citation_key(source, held)
            generated.append((source, key))
        keys[source.id] = key
    return keys, generated


def _render_tex_citations(text: str, key_by_source: dict[str, str]) -> str:
    def replace(match) -> str:
        source_ids = re.findall(r"S\d+", match.group(1))
        keys = list(dict.fromkeys(
            key_by_source[source_id] for source_id in source_ids
            if source_id in key_by_source
        ))
        return r"\cite{" + ",".join(keys) + "}"

    return _SOURCE_MARKER.sub(replace, text)


def _attach_bibliography(text: str, sidecar: Path) -> str:
    if re.search(r"\\(?:addbibresource|printbibliography)\b", text):
        resource = sidecar.name
        if resource not in text:
            point = text.find(r"\begin{document}")
            point = point if point >= 0 else 0
            text = text[:point] + f"\\addbibresource{{{resource}}}\n" + text[point:]
        return text

    bibliography = re.compile(r"\\bibliography\{([^{}]+)\}")
    match = bibliography.search(text)
    if match:
        names = [name.strip() for name in match.group(1).split(",") if name.strip()]
        if sidecar.stem not in names:
            names.append(sidecar.stem)
        return text[:match.start()] + "\\bibliography{" + ",".join(names) + "}" + text[match.end():]

    point = text.rfind(r"\end{document}")
    point = point if point >= 0 else len(text)
    style = "" if re.search(r"\\bibliographystyle\{", text) else "\\bibliographystyle{plain}\n"
    block = f"{style}\\bibliography{{{sidecar.stem}}}\n\n"
    return text[:point].rstrip() + "\n\n" + block + text[point:].lstrip()


def tex_citation_audit(path: str | Path) -> dict:
    """Report unresolved citation keys across every bibliography attached to a TeX copy."""

    path = Path(path)
    text = path.read_text(encoding="utf-8", errors="replace")
    index = _bibtex_index(_bibliography_files(path, text))
    cited = {
        key.strip() for group in re.findall(
            r"\\cite\w*(?:\[[^\]]*\])*\{([^{}]+)\}", text,
        ) for key in group.split(",") if key.strip() and key.strip() != "*"
    }
    return {
        "cited": sorted(cited), "bibliography_keys": len(index["keys"]),
        "unresolved": sorted(cited - set(index["keys"])),
        "bibliographies": [str(path) for path in _bibliography_files(path, text)],
    }


def _insert_markdown(text: str, additions: list[Addition], sources: list[EvidenceSource]) -> str:
    rendered = [(item.target_heading, render_citations(item.text, sources))
                for item in additions]
    pending = list(rendered)
    matches = list(_MD_HEADING.finditer(text))
    insertions: list[tuple[int, str]] = []
    for heading, paragraph in rendered:
        if not heading:
            continue
        match = next((held for held in matches
                      if " ".join(held.group(2).split()).casefold() == heading.casefold()), None)
        if match is None:
            continue
        level = len(match.group(1))
        end = len(text)
        for later in matches:
            if later.start() > match.start() and len(later.group(1)) <= level:
                end = later.start()
                break
        insertions.append((end, "\n\n" + paragraph + "\n"))
        pending.remove((heading, paragraph))
    for position, payload in sorted(insertions, reverse=True):
        text = text[:position].rstrip() + payload + "\n" + text[position:].lstrip()
    if pending:
        text = text.rstrip() + "\n\n## Corpus-supported context\n\n" + "\n\n".join(
            paragraph for _heading, paragraph in pending
        ) + "\n"
    used = set(_marker_ids("\n".join(item.text for item in additions)))
    references = [source for source in sources if source.id in used]
    if references:
        text = text.rstrip() + "\n\n## References added by Spiral\n\n" + "\n".join(
            f"- {reference_entry(source)}" for source in references
        ) + "\n"
    return text


def _insert_tex(path: Path, text: str, additions: list[Addition],
                sources: list[EvidenceSource]) -> str:
    """Insert grounded prose using real citation keys and a copy-owned BibTeX sidecar."""

    from spiral.documents import ensure_tex_change_color, mark_tex_change

    used = set(_marker_ids("\n".join(item.text for item in additions)))
    references = [source for source in sources if source.id in used]
    key_by_source, generated = _citation_plan(path, text, references)
    for addition in additions:
        paragraph = mark_tex_change(
            _render_tex_citations(addition.text, key_by_source)
        )
        matches = list(_TEX_SECTION.finditer(text))
        match = next((held for held in matches
                      if " ".join(held.group(1).split()).casefold()
                      == addition.target_heading.casefold()), None)
        if match:
            end = next((held.start() for held in matches if held.start() > match.start()),
                       len(text))
            text = text[:end].rstrip() + "\n\n" + paragraph + "\n\n" + text[end:].lstrip()
        else:
            point = text.rfind("\\end{document}")
            point = point if point >= 0 else len(text)
            text = (text[:point].rstrip() + "\n\n\\section*{Corpus-supported context}\n"
                    + paragraph + "\n\n" + text[point:].lstrip())
    if generated:
        sidecar = path.with_name(path.stem + ".spiral.bib")
        sidecar.write_text(
            "\n\n".join(_bibtex_entry(source, key) for source, key in generated)
            + "\n", encoding="utf-8",
        )
        text = _attach_bibliography(text, sidecar)
    return ensure_tex_change_color(text)


def append_grounded_additions(path: str | Path, kind: str,
                              result: ExpansionResult) -> None:
    path = Path(path)
    if not result.additions:
        return
    if kind == "docx":
        import docx
        from docx.shared import RGBColor
        from spiral.documents import SPIRAL_CLAY_RGB

        document = docx.Document(str(path))
        document.add_heading("Corpus-supported context", level=1)
        for addition in result.additions:
            paragraph = document.add_paragraph(
                render_citations(addition.text, result.sources)
            )
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(*SPIRAL_CLAY_RGB)
        document.add_heading("References added by Spiral", level=1)
        used = set(_marker_ids("\n".join(item.text for item in result.additions)))
        for source in result.sources:
            if source.id in used:
                paragraph = document.add_paragraph(
                    reference_entry(source), style="List Bullet"
                )
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(*SPIRAL_CLAY_RGB)
        document.save(str(path))
        return
    text = path.read_text(encoding="utf-8", errors="replace")
    if kind == "tex":
        text = _insert_tex(path, text, result.additions, result.sources)
    else:
        text = _insert_markdown(text, result.additions, result.sources)
    path.write_text(text, encoding="utf-8")


def _section_role(name: str) -> str:
    low = " ".join(name.split()).casefold()
    if "abstract" in low:
        return "abstract"
    if any(cue in low for cue in ("reference", "bibliograph")):
        return "references"
    if "acknowledg" in low:
        return "acknowledgments"
    if any(cue in low for cue in (
            "mechanism", "theoretical framework", "conceptual framework",
            "historical context", "literature review")):
        return "setup"
    if low in {"analysis", "evidence", "findings and evidence"}:
        return "results"
    try:
        from spiral.research_writer import _heading_role

        return _heading_role(name)
    except Exception:
        return "other"


def mined_structure_roles(papers, template=None) -> tuple[list[str], list[str]]:
    """Return a corpus-observed literal arc and its de-duplicated rhetorical roles."""

    try:
        from spiral.research_writer import corpus_style_profile

        profile = corpus_style_profile(papers)
        headings = list(profile.get("representative_arc") or
                        profile.get("preferred_sections") or [])
    except Exception:
        headings = []
    if not headings and template is not None:
        headings = list(getattr(template, "section_order", None) or [])
    roles = []
    for heading in headings:
        role = _section_role(str(heading))
        if role not in {"other", "appendix"} and role not in roles:
            roles.append(role)
    return headings[:10], roles[:10]


def _ordered_indices(names: list[str], target_roles: list[str]) -> list[int]:
    roles = [_section_role(name) for name in names]
    rank = {role: index for index, role in enumerate(target_roles)}
    universal = {
        "abstract": -20, "introduction": -10, "setup": 0, "methods": 10,
        "results": 20, "proof": 25, "discussion": 30, "conclusion": 40,
        "acknowledgments": 80, "references": 90, "appendix": 100,
    }
    return [index for index, _name in sorted(
        enumerate(names),
        key=lambda row: (
            -100 if roles[row[0]] == "abstract" else
            200 if roles[row[0]] == "references" else
            rank.get(roles[row[0]], 50 + universal.get(roles[row[0]], 0)),
            row[0],
        ),
    )]


def _restructure_markdown(text: str, target_roles: list[str]) -> tuple[str, bool]:
    matches = list(_MD_HEADING.finditer(text))
    if len(matches) < 3:
        return text, False
    # A single level-1 title is prefix metadata. Reorder the shallowest repeated section
    # level underneath it; nested subsections travel with their parent block.
    levels = [len(match.group(1)) for match in matches]
    candidates = [level for level, count in Counter(levels).items() if count >= 2]
    if not candidates:
        return text, False
    section_level = min(level for level in candidates
                        if not (level == 1 and levels.count(1) == 1))
    section_matches = [match for match in matches if len(match.group(1)) == section_level]
    if len(section_matches) < 2:
        return text, False
    prefix = text[:section_matches[0].start()]
    blocks = [text[match.start(): (section_matches[index + 1].start()
                                  if index + 1 < len(section_matches) else len(text))]
              for index, match in enumerate(section_matches)]
    names = [" ".join(match.group(2).split()) for match in section_matches]
    order = _ordered_indices(names, target_roles)
    if order == list(range(len(blocks))):
        return text, False
    rebuilt = prefix.rstrip() + "\n\n" + "\n\n".join(blocks[index].strip()
                                                        for index in order) + "\n"
    if Counter(block.strip() for block in blocks) != Counter(
            block.strip() for block in [blocks[index] for index in order]):
        return text, False
    return rebuilt, True


def _restructure_tex(text: str, target_roles: list[str]) -> tuple[str, bool]:
    # Appendices and bibliography material are a protected suffix. Reordering them as
    # ordinary body sections can detach proof labels or move references into the body.
    boundaries = [position for position in (
        text.find("\\appendix"), text.find("\\bibliography"),
        text.find("\\printbibliography"),
    ) if position >= 0]
    boundary = min(boundaries) if boundaries else len(text)
    matches = [match for match in _TEX_SECTION.finditer(text)
               if match.start() < boundary]
    if len(matches) < 2:
        return text, False
    prefix = text[:matches[0].start()]
    body_end = boundary
    suffix = text[body_end:]
    blocks = [text[match.start(): (matches[index + 1].start()
                                  if index + 1 < len(matches) else body_end)]
              for index, match in enumerate(matches)]
    names = [" ".join(match.group(1).split()) for match in matches]
    # A generic corpus arc cannot safely position theory-specific sections whose role
    # is unknown. Leave the native order intact and let the advisory paper audit propose
    # a reviewed reorganization instead.
    if any(_section_role(name) == "other" for name in names):
        return text, False
    order = _ordered_indices(names, target_roles)
    if order == list(range(len(blocks))):
        return text, False
    return (prefix.rstrip() + "\n\n" + "\n\n".join(
        blocks[index].strip() for index in order) + "\n\n" + suffix.lstrip(), True)


def _restructure_docx(path: Path, target_roles: list[str]) -> bool:
    import docx
    from docx.text.paragraph import Paragraph

    document = docx.Document(str(path))
    body = document._element.body
    children = list(body.iterchildren())
    starts = []
    names = []
    for index, child in enumerate(children):
        if not child.tag.endswith("}p"):
            continue
        paragraph = Paragraph(child, document._body)
        style = (paragraph.style.name if paragraph.style is not None else "") or ""
        match = re.search(r"(?i)heading\s*1\b", style)
        if match and paragraph.text.strip():
            starts.append(index)
            names.append(" ".join(paragraph.text.split()))
    if len(starts) < 2:
        return False
    prefix = children[:starts[0]]
    end = next((index for index, child in enumerate(children)
                if child.tag.endswith("}sectPr")), len(children))
    blocks = [children[start: (starts[index + 1] if index + 1 < len(starts) else end)]
              for index, start in enumerate(starts)]
    suffix = children[end:]
    order = _ordered_indices(names, target_roles)
    if order == list(range(len(blocks))):
        return False
    for child in list(body.iterchildren()):
        body.remove(child)
    for child in prefix:
        body.append(child)
    for index in order:
        for child in blocks[index]:
            body.append(child)
    for child in suffix:
        body.append(child)
    document.save(str(path))
    return True


def restructure_document(path: str | Path, kind: str, target_roles: list[str]) -> bool:
    """Reorder intact native section blocks; never regenerate section prose."""

    if len(target_roles) < 2:
        return False
    path = Path(path)
    if kind == "docx":
        return _restructure_docx(path, target_roles)
    text = path.read_text(encoding="utf-8", errors="replace")
    if kind == "tex":
        rebuilt, changed = _restructure_tex(text, target_roles)
    else:
        rebuilt, changed = _restructure_markdown(text, target_roles)
    if changed:
        path.write_text(rebuilt, encoding="utf-8")
    return changed
