"""Documents keep their structure through an edit.

The failure this guards against: flattening a paper to plain text, rewriting the lot,
and saving something that is no longer a paper — preamble gone, equations reworded,
Word styles reset. Editing happens per paragraph, through the document's own structure.
"""
import tempfile
from pathlib import Path

import pytest

from spiral.documents import (
    default_output_path, ensure_tex_change_color, mark_tex_change,
    prepare_tex_project_copy, read_document, require_distinct_output,
    write_document,
)

TEX = r"""\documentclass{article}
\usepackage{amsmath}
\begin{document}
\section{Introduction}
Additionally, this framework stands as a testament to the enduring interplay.

\begin{equation}
E = mc^2
\end{equation}

\subsection{Method}
% an internal note
We measured 42 samples \cite{smith2020}.
\end{document}"""


def _tmp(name: str, body: str = "") -> Path:
    p = Path(tempfile.mkdtemp()) / name
    if body:
        p.write_text(body)
    return p


def test_tex_protects_math_preamble_and_sectioning():
    doc = read_document(_tmp("p.tex", TEX))
    protected = "\n".join(s.text for s in doc.segments if not s.editable)
    editable = "\n".join(s.text for s in doc.editable_segments)
    for must_protect in (r"\documentclass", r"\usepackage", r"\begin{document}",
                         r"\section{Introduction}", r"E = mc^2", r"\subsection{Method}",
                         "% an internal note", r"\end{document}"):
        assert must_protect in protected, f"{must_protect} was left editable"
        assert must_protect not in editable
    assert "stands as a testament" in editable


def test_tex_edit_replaces_only_the_targeted_paragraph():
    src = _tmp("p.tex", TEX)
    doc = read_document(src)
    target = doc.editable_segments[0]
    out = write_document(doc, {target.index: "This framework links theory and practice."},
                         src.with_name("p.edited.tex"))
    edited = out.read_text()
    assert "This framework links theory and practice." in edited
    assert r"\definecolor{SpiralClay}{HTML}{D97757}" in edited
    assert r"{\color{SpiralClay}This framework links theory and practice.}" in edited
    assert "stands as a testament" not in edited
    assert r"E = mc^2" in edited and r"\documentclass{article}" in edited


def test_tex_writer_preserves_every_untouched_byte_around_multiline_commands(tmp_path):
    raw = (
        r"\documentclass{article}" "\n\n"
        r"\begin{document}" "\n\n"
        "\\title{A title split across\n"
        "two source lines}\n\n"
        r"\maketitle" "\n\n"
        "Additionally, this sentence can change.\n\n"
        r"\end{document}" "\n"
    )
    source = tmp_path / "main.tex"
    source.write_text(raw)
    document = read_document(source)
    target = next(segment for segment in document.editable_segments
                  if "Additionally" in segment.text)
    output = write_document(
        document, {target.index: "This sentence changed."},
        tmp_path / "main.edited.tex",
    )
    expected = ensure_tex_change_color(raw.replace(
        "Additionally, this sentence can change.",
        mark_tex_change("This sentence changed."),
    ))
    assert output.read_text() == expected
    assert "across\ntwo" in output.read_text()
    assert "across\n\ntwo" not in output.read_text()


def test_original_is_never_modified():
    src = _tmp("p.tex", TEX)
    before = src.read_text()
    doc = read_document(src)
    write_document(doc, {doc.editable_segments[0].index: "replaced"},
                   src.with_name("p.edited.tex"))
    assert src.read_text() == before


def test_writer_rejects_source_path_symlink_and_hardlink(tmp_path):
    src = tmp_path / "paper.tex"
    src.write_text(TEX)
    doc = read_document(src)
    replacement = {doc.editable_segments[0].index: "replacement"}
    with pytest.raises(ValueError, match="cannot be the source"):
        write_document(doc, replacement, src)

    symlink = tmp_path / "alias.tex"
    symlink.symlink_to(src)
    with pytest.raises(ValueError, match="cannot be the source"):
        write_document(doc, replacement, symlink)

    hardlink = tmp_path / "hardlink.tex"
    hardlink.hardlink_to(src)
    with pytest.raises(ValueError, match="aliases the source"):
        require_distinct_output(src, hardlink)
    assert src.read_text() == TEX


def test_tex_project_copy_preserves_sources_but_omits_dependency_caches(tmp_path):
    project = tmp_path / "paper"
    project.mkdir()
    main = project / "main.tex"
    main.write_text(TEX + "\n\\input{sections/result.tex}\n")
    (project / "references.bib").write_text("@article{x,title={X}}\n")
    (project / "sections").mkdir()
    (project / "sections/result.tex").write_text("Result text.\n")
    (project / "figures").mkdir()
    (project / "figures/data.dat").write_text("x y\n1 2\n")
    (project / ".spiral").mkdir()
    (project / ".spiral/cache").write_text("large cache")
    (project / "lean/.lake").mkdir(parents=True)
    (project / "lean/.lake/dependency").write_text("large dependency")
    before = main.read_bytes()

    copied_main = prepare_tex_project_copy(main)

    assert copied_main.parent.name == "paper_spiral"
    assert copied_main.read_bytes() == before
    assert (copied_main.parent / "references.bib").is_file()
    assert (copied_main.parent / "sections/result.tex").is_file()
    assert (copied_main.parent / "figures/data.dat").is_file()
    assert not (copied_main.parent / ".spiral").exists()
    assert not (copied_main.parent / "lean/.lake").exists()
    assert main.read_bytes() == before


def test_markdown_headings_are_not_rewritten():
    doc = read_document(_tmp("n.md", "# Title\n\nSome prose here that can change.\n"))
    assert any(s.kind == "heading" and not s.editable for s in doc.segments)
    assert any("Some prose" in s.text for s in doc.editable_segments)


def test_default_output_never_overwrites_the_source():
    assert default_output_path(Path("a/paper.tex")).name == "paper_spiral.tex"
    assert default_output_path(Path("a/report.docx")).name == "report_spiral.docx"
    # a PDF cannot be faithfully rewritten, so its edit is emitted as markdown
    assert default_output_path(Path("a/scan.pdf")).name == "scan_spiral.md"


def test_with_replacements_does_not_mutate_the_document():
    doc = read_document(_tmp("p.tex", TEX))
    idx = doc.editable_segments[0].index
    once = doc.with_replacements({idx: "NEW"})
    assert "NEW" in once
    assert "NEW" not in doc.text()          # pure function


# ------------------------------------------------------------------ docx
docx = pytest.importorskip("docx")


def test_docx_keeps_headings_and_styles_through_an_edit():
    path = _tmp("r.docx")
    d = docx.Document()
    d.add_heading("Results", level=1)
    d.add_paragraph("Additionally, this stands as a testament to the robust interplay.")
    d.save(str(path))

    doc = read_document(path)
    assert not doc.segments[0].editable, "a heading must not be rewritten"
    assert doc.segments[1].editable

    out = write_document(doc, {1: "This connects several ideas."},
                         path.with_name("r.edited.docx"))
    back = docx.Document(str(out))
    assert back.paragraphs[0].text == "Results"
    assert back.paragraphs[0].style.name.lower().startswith("heading")
    assert "testament" not in back.paragraphs[1].text
    assert "This connects several ideas." == back.paragraphs[1].text
    assert str(back.paragraphs[1].runs[0].font.color.rgb) == "D97757"


def test_docx_protects_captions_and_review_markup():
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement

    path = _tmp("reviewed.docx")
    backing = docx.Document()
    backing.styles.add_style("Image Caption", WD_STYLE_TYPE.PARAGRAPH)
    backing.add_paragraph("Figure 1. A measured response.", style="Image Caption")
    reviewed = backing.add_paragraph("Additionally, the measured response changed.")
    reviewed._p.append(OxmlElement("w:ins"))
    backing.add_paragraph("Additionally, this plain paragraph can be edited.")
    backing.save(str(path))

    parsed = read_document(path)
    assert parsed.segments[0].editable is False
    assert parsed.segments[1].editable is False
    assert parsed.segments[2].editable is True


def test_prose_plan_lists_only_the_stages_that_will_run():
    """A plan showing work that gets skipped is a plan you stop trusting."""
    from spiral.prose_ui import prose_plan

    measure_only = prose_plan(with_corpus=False, rewriting=False)
    assert [m.title for m in measure_only.milestones] == ["read the document"]

    full = prose_plan(with_corpus=True, rewriting=True)
    titles = [m.title for m in full.milestones]
    assert titles == ["read the document", "field template", "rewrite", "write out"]
    assert full.task_count == 9
