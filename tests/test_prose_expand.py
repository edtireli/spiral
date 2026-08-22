"""Grounded expansion and deterministic corpus-structure alignment."""
from __future__ import annotations

import json
import re
from pathlib import Path
from types import SimpleNamespace

from spiral.prose_expand import (
    Addition, EvidenceAnchor, EvidenceSource, ExpansionResult,
    _anchor_candidates, append_grounded_additions, build_evidence_packet,
    generate_grounded_additions,
    restructure_document, tex_citation_audit, validate_additions,
    validate_entailment_audit,
)
from spiral.research_corpus import Paper


def _source(identifier="close", relevance=0.6):
    abstract = (
        "We measured surface temperatures across 30 neighborhoods and found that "
        "streets with denser tree canopy were 2 C cooler during afternoon observations."
    )
    paper = Paper(
        arxiv_id=identifier, title="Street tree canopy and afternoon heat",
        authors=["Ada Green", "Bo Shade"], published="2024-05-01",
        abstract=abstract, text=("Results. " + abstract + " ") * 30,
        body_source="jats", source="pmc", doi="10.1000/tree.1",
    )
    return paper, {"id": paper.bare_id, "relevance": relevance,
                   "matched_terminology": ["tree canopy"]}


def _packet():
    return [EvidenceSource(
        id="S1", paper_id="close", title="Street tree canopy and afternoon heat",
        authors=["Ada Green", "Bo Shade"], year="2024", doi="10.1000/tree.1",
        relevance=0.6,
        anchors=[EvidenceAnchor(
            "S1-A1",
            "We measured surface temperatures across 30 neighborhoods and found that "
            "streets with denser tree canopy were 2 C cooler during afternoon observations.",
        )],
    )]


def test_evidence_packet_keeps_every_source_that_passed_deep_relevance_gates():
    close, close_row = _source("close", 0.6)
    weak, weak_row = _source("weak", 0.26)
    packet = build_evidence_packet(
        [close, weak], [close_row, weak_row],
        "Urban tree canopy reduces afternoon heat in neighborhoods.",
        minimum_sources=1,
    )
    assert [source.paper_id for source in packet] == ["close", "weak"]
    assert packet[0].anchors and packet[0].anchors[0].text in close.text


def test_abstract_contribution_outranks_equation_and_section_navigation():
    paper, _row = _source()
    paper.abstract = (
        "We derive Carrollian amplitudes and show that the transform has a regular "
        "positive-energy representation for MHV scattering."
    )
    paper.text = (
        "See section 4 where we discuss Carrollian amplitudes. "
        + (r"\begin{equation} Carrollian amplitudes positive energy MHV transform "
           r"x=y+z \end{equation}. " * 30)
    )
    anchors = _anchor_candidates(
        paper, {"carrollian", "amplitudes", "positive-energy", "mhv", "transform"},
    )
    assert anchors[0][1].startswith("We derive Carrollian amplitudes")
    assert all("See section 4" not in text for _score, text in anchors)


def test_self_contained_abstract_claim_outranks_deictic_fragment():
    paper, _row = _source()
    paper.abstract = (
        "This result is generalized to all loop orders. "
        "We compute one-loop four-point Carrollian MHV amplitudes as differential "
        "operators acting on tree-level amplitudes."
    )
    anchors = _anchor_candidates(
        paper, {"carrollian", "mhv", "amplitudes", "loop", "operators"},
    )
    assert anchors[0][1].startswith("We compute one-loop four-point")


def test_evidence_packet_refuses_a_two_source_corpus_by_default():
    first, first_row = _source("first", 0.6)
    second, second_row = _source("second", 0.5)
    assert build_evidence_packet(
        [first, second], [first_row, second_row],
        "Urban tree canopy reduces afternoon heat in neighborhoods.",
    ) == []


def test_addition_gate_requires_citations_exact_anchors_and_supported_numbers():
    sources = _packet()
    valid = {"additions": [{
        "target_heading": "Results",
        "text": ("Across 30 neighborhoods, streets with denser tree canopy had afternoon "
                 "temperatures that were 2 C lower [S1]."),
        "evidence": [{"source": "S1", "anchor": "S1-A1"}],
    }]}
    accepted = validate_additions(
        valid, sources, "# Trees\n\n## Results\n\nOriginal.", ["Trees", "Results"],
        max_added_words=100,
    )
    assert len(accepted.additions) == 1

    uncited = json.loads(json.dumps(valid))
    uncited["additions"][0]["text"] = uncited["additions"][0]["text"].replace(
        " [S1]", ""
    )
    assert "no source marker" in validate_additions(
        uncited, sources, "draft", ["Results"], max_added_words=100,
    ).issues[0]

    invented = json.loads(json.dumps(valid))
    invented["additions"][0]["text"] = invented["additions"][0]["text"].replace(
        "30", "300"
    )
    assert any("number absent" in issue for issue in validate_additions(
        invented, sources, "draft", ["Results"], max_added_words=100,
    ).issues)

    wrong_anchor = json.loads(json.dumps(valid))
    wrong_anchor["additions"][0]["evidence"][0]["anchor"] = "S1-A99"
    assert any("lacks an exact anchor" in issue for issue in validate_additions(
        wrong_anchor, sources, "draft", ["Results"], max_added_words=100,
    ).issues)

    anchor_inline = {"additions": [{
        "target_heading": "Results",
        "text": ("Across 30 neighborhoods, denser tree canopy was associated with "
                 "afternoon temperatures that were 2 C lower [S1-A1]."),
    }]}
    derived = validate_additions(
        anchor_inline, sources, "draft", ["Results"], max_added_words=100,
    )
    assert len(derived.additions) == 1
    assert derived.additions[0].text.endswith("[S1].")
    assert derived.additions[0].evidence == [{"source": "S1", "anchor": "S1-A1"}]


def test_entailment_audit_must_cover_every_claim_with_allowed_anchor():
    claims = [{"claim_id": "C1", "allowed_anchors": ["S1-A1"]}]
    assert validate_entailment_audit(
        {"claims": [{"claim_id": "C1", "supported": True,
                     "anchor_ids": ["S1-A1"]}]}, claims,
    ) == []
    assert validate_entailment_audit(
        {"claims": [{"claim_id": "C1", "supported": True,
                     "anchor_ids": ["S2-A1"]}]}, claims,
    ) == ["claim C1 used an unverified anchor"]


class _LLM:
    def __init__(self, replies):
        self.replies = list(replies)

    def chat(self, *args, **kwargs):
        return SimpleNamespace(text=json.dumps(self.replies.pop(0)))


class _Cfg:
    critic = SimpleNamespace(name="critic")
    keep_alive = "5m"

    @staticmethod
    def spec_for(model):
        return SimpleNamespace(num_ctx=8192)


def test_generation_uses_independent_entailment_audit():
    paper, row = _source()
    proposal = {"additions": [{
        "target_heading": "Introduction",
        "text": ("Across 30 neighborhoods, denser tree canopy was associated with "
                 "afternoon temperatures that were 2 C lower [S1]."),
        "evidence": [{"source": "S1", "anchor": "S1-A1"}],
    }]}
    audit = {"claims": [{"claim_id": "C1", "supported": True,
                         "anchor_ids": ["S1-A1"]}]}
    result = generate_grounded_additions(
        _LLM([proposal, audit]), _Cfg(), "writer",
        "# Trees\n\n## Introduction\n\nTree canopy affects urban heat.", "markdown",
        [paper], [row], rounds=1, minimum_sources=1,
    )
    assert result.added_words > 0


def test_generation_keeps_supported_paragraph_and_drops_failed_paragraph():
    paper, row = _source()
    proposal = {"additions": [
        {
            "target_heading": "Introduction",
            "text": ("Across 30 neighborhoods, denser tree canopy was associated with "
                     "afternoon temperatures that were 2 C lower [S1-A1]."),
        },
        {
            "target_heading": "Introduction",
            "text": ("The same 30 neighborhoods prove that tree canopy always lowers "
                     "afternoon temperatures by 2 C [S1-A1]."),
        },
    ]}
    audit = {"claims": [
        {"claim_id": "C1", "supported": True, "anchor_ids": ["S1-A1"]},
        {"claim_id": "C2", "supported": False, "anchor_ids": ["S1-A1"]},
    ]}
    result = generate_grounded_additions(
        _LLM([proposal, audit]), _Cfg(), "writer",
        "# Trees\n\n## Introduction\n\nTree canopy affects urban heat.", "markdown",
        [paper], [row], rounds=1, minimum_sources=1,
    )
    assert len(result.additions) == 1
    assert "associated" in result.additions[0].text
    assert any("not entailed" in issue for issue in result.issues)


def test_generation_requires_four_distinct_cited_sources_by_default():
    papers = []
    rows = []
    for index in range(4):
        paper, row = _source(f"source-{index}", 0.6 - index * 0.02)
        papers.append(paper)
        rows.append(row)
    proposal = {"additions": [{
        "target_heading": "Introduction",
        "text": ("Across 30 neighborhoods, denser tree canopy was associated with "
                 "afternoon temperatures that were 2 C lower [S1-A1; S2-A1; S3-A1]."),
    }]}
    result = generate_grounded_additions(
        _LLM([proposal]), _Cfg(), "writer",
        "# Trees\n\n## Introduction\n\nTree canopy affects urban heat.", "markdown",
        papers, rows, rounds=1,
    )
    assert result.additions == []
    assert any("use at least 4" in issue for issue in result.issues)


def test_generation_accepts_any_four_close_sources_not_only_first_four():
    papers = []
    rows = []
    for index in range(5):
        paper, row = _source(f"source-{index}", 0.7 - index * 0.02)
        papers.append(paper)
        rows.append(row)
    proposal = {"additions": [{
        "target_heading": "Introduction",
        "text": (
            "Across 30 neighborhoods, denser tree canopy was associated with afternoon "
            "temperatures that were 2 C lower [S2-A1; S3-A1; S4-A1; S5-A1]."
        ),
    }]}
    audit = {"claims": [{
        "claim_id": "C1", "supported": True,
        "anchor_ids": ["S2-A1", "S3-A1", "S4-A1", "S5-A1"],
    }]}
    result = generate_grounded_additions(
        _LLM([proposal, audit]), _Cfg(), "writer",
        "# Trees\n\n## Introduction\n\nTree canopy affects urban heat.", "markdown",
        papers, rows, rounds=1,
    )
    assert len(result.additions) == 1
    assert {row["source"] for row in result.additions[0].evidence} == {
        "S2", "S3", "S4", "S5",
    }


def test_generation_falls_back_to_independently_gated_source_sentences():
    papers = []
    rows = []
    for index in range(4):
        paper, row = _source(f"source-{index}", 0.7 - index * 0.02)
        papers.append(paper)
        rows.append(row)
    broad_but_incomplete = {"additions": [{
        "target_heading": "Introduction",
        "text": ("Across 30 neighborhoods, denser tree canopy was associated with "
                 "afternoon temperatures that were 2 C lower [S1-A1]."),
    }]}
    individual = [{"additions": [{
        "target_heading": "Introduction",
        "text": ("Across 30 neighborhoods, denser tree canopy was associated with "
                 f"afternoon temperatures that were 2 C lower [S{index}-A1]."),
    }]} for index in range(1, 5)]
    audit = {"claims": [
        {"claim_id": f"C{index}", "supported": True,
         "anchor_ids": [f"S{index}-A1"]}
        for index in range(1, 5)
    ]}
    result = generate_grounded_additions(
        _LLM([broad_but_incomplete, *individual, audit]), _Cfg(), "writer",
        "# Trees\n\n## Introduction\n\nTree canopy affects urban heat.", "markdown",
        papers, rows, rounds=1,
    )
    assert len(result.additions) == 4
    assert {source.id for source in result.sources} == {"S1", "S2", "S3", "S4"}


def test_tex_generation_rejects_unicode_extraction_artifacts():
    paper, row = _source()
    proposal = {"additions": [{
        "target_heading": "Introduction",
        "text": ("Across 30 neighborhoods, the measured afternoon temperature change "
                 "for streets with denser tree canopy was Δ = 2 C [S1-A1]."),
    }]}
    result = generate_grounded_additions(
        _LLM([proposal]), _Cfg(), "writer",
        r"\section{Introduction}" "\nTree canopy affects urban heat.", "tex",
        [paper], [row], rounds=1, minimum_sources=1,
    )
    assert result.additions == []
    assert any("non-ASCII" in issue for issue in result.issues)


def test_markdown_additions_render_academic_citation_and_reference(tmp_path):
    path = tmp_path / "draft.md"
    path.write_text("# Trees\n\n## Results\n\nOriginal result.\n\n## Discussion\n\nOriginal discussion.\n")
    result = ExpansionResult(
        additions=[Addition(
            "Results", "Denser canopy was associated with lower afternoon heat [S1].",
            [{"source": "S1", "anchor": "S1-A1"}],
        )], sources=_packet(),
    )
    append_grounded_additions(path, "markdown", result)
    rendered = path.read_text()
    assert "(Green et al., 2024)" in rendered
    assert rendered.index("Denser canopy") < rendered.index("## Discussion")
    assert "## References added by Spiral" in rendered
    assert "https://doi.org/10.1000/tree.1" in rendered


def test_docx_additions_use_spiral_clay_for_review(tmp_path):
    docx = __import__("docx")
    path = tmp_path / "draft.docx"
    document = docx.Document()
    document.add_heading("Introduction", level=1)
    document.add_paragraph("Original text.")
    document.save(path)
    result = ExpansionResult(
        additions=[Addition(
            "Introduction", "Denser canopy was associated with lower heat [S1].",
            [{"source": "S1", "anchor": "S1-A1"}],
        )], sources=_packet(),
    )
    append_grounded_additions(path, "docx", result)
    reopened = docx.Document(path)
    changed = next(paragraph for paragraph in reopened.paragraphs
                   if "Denser canopy" in paragraph.text)
    assert str(changed.runs[0].font.color.rgb) == "D97757"


def test_pubmed_surname_initials_render_as_surname_not_initials(tmp_path):
    path = tmp_path / "draft.md"
    path.write_text("# Trees\n\n## Introduction\n\nOriginal.\n")
    source = _packet()[0]
    source.authors = ["Ettinger AK", "Bratman GN"]
    result = ExpansionResult(
        additions=[Addition(
            "Introduction", "Canopy was associated with afternoon cooling [S1].",
            [{"source": "S1", "anchor": "S1-A1"}],
        )], sources=[source],
    )
    append_grounded_additions(path, "markdown", result)
    assert "(Ettinger et al., 2024)" in path.read_text()


def test_tex_beef_up_reuses_held_bibtex_key_without_mutating_database(tmp_path):
    path = tmp_path / "paper.edited.tex"
    path.write_text(
        r"\documentclass{article}" "\n"
        r"\begin{document}" "\n"
        r"\section{Introduction}" "\nOriginal.\n"
        r"\bibliographystyle{plain}" "\n"
        r"\bibliography{references}" "\n"
        r"\end{document}" "\n"
    )
    bibliography = tmp_path / "references.bib"
    bibliography.write_text(
        "@article{HeldTree2024,\n"
        "  title = {Street tree canopy and afternoon heat},\n"
        "  doi = {10.1000/tree.1},\n"
        "  year = {2024}\n}\n"
    )
    before = bibliography.read_bytes()
    result = ExpansionResult(
        additions=[Addition(
            "Introduction", "Denser canopy was associated with lower heat [S1].",
            [{"source": "S1", "anchor": "S1-A1"}],
        )], sources=_packet(),
    )

    append_grounded_additions(path, "tex", result)

    rendered = path.read_text()
    assert r"\cite{HeldTree2024}" in rendered
    assert r"\definecolor{SpiralClay}{HTML}{D97757}" in rendered
    assert r"{\color{SpiralClay}" in rendered
    assert "References added by Spiral" not in rendered
    assert r"\bibliography{references}" in rendered
    assert bibliography.read_bytes() == before
    assert not (tmp_path / "paper.edited.spiral.bib").exists()
    assert tex_citation_audit(path)["unresolved"] == []


def test_tex_beef_up_writes_copy_owned_bibtex_sidecar_and_resolves_citation(tmp_path):
    path = tmp_path / "paper.edited.tex"
    path.write_text(
        r"\documentclass{article}" "\n"
        r"\begin{document}" "\n"
        r"\section{Introduction}" "\nOriginal.\n"
        r"\bibliographystyle{plain}" "\n"
        r"\bibliography{references}" "\n"
        r"\end{document}" "\n"
    )
    bibliography = tmp_path / "references.bib"
    bibliography.write_text("@article{Existing, title={Another paper}, year={2020}}\n")
    before = bibliography.read_bytes()
    source = _packet()[0]
    source.doi = "10.1000/new-tree"
    source.title = "New evidence on tree canopy"
    result = ExpansionResult(
        additions=[Addition(
            "Introduction", "Denser canopy was associated with lower heat [S1].",
            [{"source": "S1", "anchor": "S1-A1"}],
        )], sources=[source],
    )

    append_grounded_additions(path, "tex", result)

    rendered = path.read_text()
    sidecar = tmp_path / "paper.edited.spiral.bib"
    assert sidecar.is_file()
    assert "10.1000/new-tree" in sidecar.read_text()
    assert re.search(r"\\cite\{Spiral[^}]+\}", rendered)
    assert r"\bibliography{references,paper.edited.spiral}" in rendered
    assert bibliography.read_bytes() == before
    assert tex_citation_audit(path)["unresolved"] == []


def test_markdown_restructure_moves_intact_blocks_toward_corpus_roles(tmp_path):
    path = tmp_path / "draft.md"
    path.write_text(
        "# Study\n\n## Discussion\n\nDISCUSSION TOKEN\n\n"
        "## Methods\n\nMETHOD TOKEN\n\n## Results\n\nRESULT TOKEN\n"
    )
    assert restructure_document(
        path, "markdown", ["introduction", "methods", "results", "discussion"],
    ) is True
    text = path.read_text()
    assert text.index("METHOD TOKEN") < text.index("RESULT TOKEN") < text.index(
        "DISCUSSION TOKEN"
    )
    assert all(text.count(token) == 1 for token in
               ("METHOD TOKEN", "RESULT TOKEN", "DISCUSSION TOKEN"))


def test_tex_restructure_keeps_appendix_and_bibliography_as_fixed_suffix(tmp_path):
    path = tmp_path / "paper.edited.tex"
    path.write_text(
        r"\documentclass{article}" "\n"
        r"\begin{document}" "\n"
        r"\section{Discussion}" "\nDISCUSSION\n"
        r"\section{Methods}" "\nMETHODS\n"
        r"\section{Results}" "\nRESULTS\n"
        r"\appendix" "\n"
        r"\section{Proof certificate}" "\nAPPENDIX TOKEN\n"
        r"\bibliography{references}" "\n"
        r"\end{document}" "\n"
    )
    assert restructure_document(
        path, "tex", ["methods", "results", "discussion"],
    )
    text = path.read_text()
    assert text.index("METHODS") < text.index("RESULTS") < text.index("DISCUSSION")
    assert text.index("DISCUSSION") < text.index(r"\appendix")
    assert text.index("APPENDIX TOKEN") < text.index(r"\bibliography")
    assert text.count("APPENDIX TOKEN") == 1


def test_tex_restructure_does_not_guess_the_role_of_theory_specific_sections(tmp_path):
    path = tmp_path / "paper.edited.tex"
    original = (
        r"\documentclass{article}" "\n"
        r"\begin{document}" "\n"
        r"\section{Introduction}" "\nINTRO\n"
        r"\section{Positive-energy transform}" "\nTHEORY\n"
        r"\section{Main theorem}" "\nTHEOREM\n"
        r"\section{Discussion}" "\nDISCUSSION\n"
        r"\end{document}" "\n"
    )
    path.write_text(original)
    assert restructure_document(
        path, "tex", ["introduction", "setup", "results", "discussion"],
    ) is False
    assert path.read_text() == original


def test_mechanism_section_is_setup_and_stays_before_discussion(tmp_path):
    path = tmp_path / "draft.md"
    path.write_text(
        "# Study\n\n## Discussion\n\nDISCUSSION\n\n"
        "## Mechanisms and Evidence\n\nMECHANISMS\n\n## Conclusion\n\nEND\n"
    )
    assert restructure_document(
        path, "markdown", ["introduction", "setup", "results", "discussion"],
    )
    text = path.read_text()
    assert text.index("MECHANISMS") < text.index("DISCUSSION") < text.index("END")


def test_docx_restructure_moves_tables_with_their_section(tmp_path):
    import docx

    path = tmp_path / "draft.docx"
    document = docx.Document()
    document.add_heading("Discussion", level=1)
    document.add_paragraph("DISCUSSION TOKEN")
    document.add_heading("Methods", level=1)
    document.add_paragraph("METHOD TOKEN")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "METHOD TABLE TOKEN"
    document.add_heading("Results", level=1)
    document.add_paragraph("RESULT TOKEN")
    document.save(path)

    assert restructure_document(path, "docx", ["methods", "results", "discussion"])
    reopened = docx.Document(path)
    body_text = " ".join(node.text or "" for node in reopened._element.body.iter())
    assert body_text.index("METHOD TOKEN") < body_text.index("METHOD TABLE TOKEN")
    assert body_text.index("METHOD TABLE TOKEN") < body_text.index("RESULT TOKEN")
    assert body_text.index("RESULT TOKEN") < body_text.index("DISCUSSION TOKEN")
